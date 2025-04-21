import time
import sys
from sys import executable
from datetime import date, datetime
from meet_functions import *
from google_calendar import next_event_details
from subprocess import Popen
import os
import signal
from docx import Document
from assemblyapi import get_transcript

# function to get current time
def get_time():
    now = datetime.now()
    current_time = now.strftime("%H:%M:%S")
    return str(current_time)
def save_transcript_as_doc(transcript_text):
    now = datetime.now().strftime("%Y-%m-%d %H-%M-%S")
    doc = Document()
    doc.add_heading('Google Meet Transcript', 0)
    doc.add_paragraph(f"Generated on: {now}\n")
    doc.add_paragraph(transcript_text)

    filename = f"transcript_{now}.docx"
    filepath = os.path.join(os.getcwd(), filename)
    doc.save(filepath)
    print(f"[INFO] Transcript saved at: {filepath}")
if __name__ == "__main__":
    print("[INFO] Fetching next event details from Google Calendar...")
    event_dic = next_event_details()

    if event_dic is None:
        print("[WARN] No upcoming events found. Exiting.")
        sys.exit()

    print(f"[DEBUG] Event details fetched: {event_dic}")

    if not(event_dic['start_time'][:10].strip('T') == str(date.today())):
        print("[INFO] No meetings scheduled for today. Exiting.")
        sys.exit()
    
    print("[INFO] Waiting for meeting start time...")
    while True:
        now_time = get_time()
        if event_dic['start_time'][11:19].strip('T') <= now_time:
            print(f"[INFO] Meeting time reached: {now_time}. Proceeding to join.")
            break
        else:
            print(f"[WAIT] Current time: {now_time} | Scheduled start: {event_dic['start_time'][11:19]}")
            time.sleep(30)
    
    meet_code = event_dic['link'][24:]
    print(f"[INFO] Google Meet code extracted: {meet_code}")

    print("[INFO] Initializing browser and joining meeting...")
    driver = initial_stuff(meet_code)

    chat_dic = {}
    new_msg = []
    prev_chat = {}
    flag = True

    print("[INFO] Starting audio recording subprocess...")
    p = Popen([executable, 'C:/Users/Ankita ghosh/Desktop/Zense-Project/record.py'], shell=False)
    processId = p.pid
    print(f"[DEBUG] Audio recording process ID: {processId}")

    print("[INFO] Meeting monitoring started...")
    while True:
        new_msg = scrape(driver, chat_dic)

        now_time = get_time()
        if event_dic['end_time'][11:19].strip('T') <= now_time and event_dic['end_time'][:10].strip('T') == str(date.today()):
            print("[INFO] Scheduled meeting end time reached.")
            break
        
        if not check_meeting(driver):
            print("[INFO] Detected meeting end by UI state.")
            break

        time.sleep(10)  # Reduce load on CPU


    print("[INFO] Meeting has ended. Proceeding to cleanup...")

    close_driver(driver)
  
    with open("stop_recording.flag", "w") as f:
        f.write("stop")
    p.wait()  # Wait for process to end

    # Clean up flag file
    os.remove("stop_recording.flag")
    time.sleep(1)
    print("[INFO] Audio recording stopped.")

    print("[INFO] Requesting transcript and summary from AssemblyAI...")
    content = get_transcript()
    print("[DEBUG] Transcription and summary received.")

    save_transcript_as_doc(content)

    sys.exit()
